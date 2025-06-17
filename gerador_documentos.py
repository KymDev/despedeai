"""
Gerador de Documentos do Kit Rescisão Premium
Módulo para geração de PDF, DOCX e XLSX personalizados
"""

import os
import zipfile
from datetime import datetime
from decimal import Decimal
from io import BytesIO
from weasyprint import HTML
from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from openpyxl.utils import get_column_letter

class GeradorDocumentos:
    def __init__(self):
        self.templates_path = os.path.join(os.path.dirname(__file__), 'templates_premium')
        os.makedirs(self.templates_path, exist_ok=True)
    
    def gerar_relatorio_pdf(self, dados_calculo, dados_usuario):
        """Gera relatório personalizado em PDF"""
        
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <style>
                body {{ font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; }}
                .header {{ text-align: center; margin-bottom: 30px; }}
                .logo {{ font-size: 24px; font-weight: bold; color: #2c3e50; }}
                .subtitle {{ color: #7f8c8d; margin-top: 5px; }}
                .section {{ margin: 20px 0; }}
                .section h2 {{ color: #2c3e50; border-bottom: 2px solid #3498db; padding-bottom: 5px; }}
                .table {{ width: 100%; border-collapse: collapse; margin: 15px 0; }}
                .table th, .table td {{ border: 1px solid #ddd; padding: 12px; text-align: left; }}
                .table th {{ background-color: #3498db; color: white; }}
                .table tr:nth-child(even) {{ background-color: #f2f2f2; }}
                .highlight {{ background-color: #e8f5e8; font-weight: bold; }}
                .warning {{ background-color: #fff3cd; padding: 10px; border-left: 4px solid #ffc107; }}
                .footer {{ margin-top: 40px; text-align: center; color: #7f8c8d; font-size: 12px; }}
            </style>
        </head>
        <body>
            <div class="header">
                <div class="logo">DESPEDE AI</div>
                <div class="subtitle">Relatório Personalizado de Rescisão Trabalhista</div>
                <div class="subtitle">Gerado em {datetime.now().strftime('%d/%m/%Y às %H:%M')}</div>
            </div>
            
            <div class="section">
                <h2>📋 Dados do Funcionário</h2>
                <table class="table">
                    <tr><td><strong>Nome:</strong></td><td>{dados_usuario.get('nome', 'N/A')}</td></tr>
                    <tr><td><strong>Cargo:</strong></td><td>{dados_usuario.get('cargo', 'N/A')}</td></tr>
                    <tr><td><strong>Empresa:</strong></td><td>{dados_usuario.get('empresa', 'N/A')}</td></tr>
                    <tr><td><strong>Salário:</strong></td><td>R$ {dados_usuario.get('salario', 0):,.2f}</td></tr>
                    <tr><td><strong>Data de Admissão:</strong></td><td>{dados_usuario.get('data_admissao', 'N/A')}</td></tr>
                    <tr><td><strong>Data de Demissão:</strong></td><td>{dados_usuario.get('data_demissao', 'N/A')}</td></tr>
                    <tr><td><strong>Tipo de Demissão:</strong></td><td>{self._formatar_tipo_demissao(dados_usuario.get('tipo_demissao', 'N/A'))}</td></tr>
                </table>
            </div>
            
            <div class="section">
                <h2>💰 Cálculo Detalhado das Verbas</h2>
                <table class="table">
                    <thead>
                        <tr>
                            <th>Verba</th>
                            <th>Cálculo</th>
                            <th>Valor (R$)</th>
                        </tr>
                    </thead>
                    <tbody>
                        {self._gerar_linhas_verbas(dados_calculo)}
                    </tbody>
                </table>
            </div>
            
            <div class="section">
                <h2>📊 Resumo Financeiro</h2>
                <table class="table">
                    <tr><td><strong>Total Bruto:</strong></td><td class="highlight">R$ {dados_calculo['resumo']['total_bruto']:,.2f}</td></tr>
                    <tr><td><strong>(-) INSS:</strong></td><td>R$ {dados_calculo['resumo']['desconto_inss']:,.2f}</td></tr>
                    <tr><td><strong>(-) IRRF:</strong></td><td>R$ {dados_calculo['resumo']['desconto_irrf']:,.2f}</td></tr>
                    <tr><td><strong>Total Líquido:</strong></td><td class="highlight">R$ {dados_calculo['resumo']['total_liquido']:,.2f}</td></tr>
                </table>
            </div>
            
            <div class="section">
                <h2>💡 Estratégias e Recomendações</h2>
                {self._gerar_recomendacoes(dados_calculo, dados_usuario)}
            </div>
            
            <div class="warning">
                <strong>⚠️ Importante:</strong> Este relatório é uma estimativa baseada nas informações fornecidas. 
                Consulte sempre um contador ou advogado trabalhista para validação oficial dos cálculos.
            </div>
            
            <div class="footer">
                <p>Relatório gerado por DESPEDE AI - www.despedeai.site</p>
                <p>© 2025 - Todos os direitos reservados</p>
            </div>
        </body>
        </html>
        """
        
        # Gera PDF
        pdf_buffer = BytesIO()
        HTML(string=html_content).write_pdf(pdf_buffer)
        pdf_buffer.seek(0)
        
        return pdf_buffer
    
    def gerar_modelos_docx(self, dados_usuario):
        """Gera modelos DOCX editáveis"""
        documentos = {}
        
        # 1. Recibo de Quitação
        doc_recibo = Document()
        doc_recibo.add_heading('RECIBO DE QUITAÇÃO DE VERBAS RESCISÓRIAS', 0)
        
        p = doc_recibo.add_paragraph()
        p.add_run('Eu, ').bold = False
        p.add_run(f"{dados_usuario.get('nome', '[NOME]')}").bold = True
        p.add_run(', portador do CPF nº ').bold = False
        p.add_run('[CPF]').bold = True
        p.add_run(f', ex-funcionário da empresa ').bold = False
        p.add_run(f"{dados_usuario.get('empresa', '[EMPRESA]')}").bold = True
        p.add_run(f', onde exercia o cargo de ').bold = False
        p.add_run(f"{dados_usuario.get('cargo', '[CARGO]')}").bold = True
        p.add_run(', declaro ter recebido todas as verbas rescisórias a que tinha direito.').bold = False
        
        doc_recibo.add_paragraph('\\nVerbas recebidas:')
        doc_recibo.add_paragraph('• Saldo de salário: R$ [VALOR]')
        doc_recibo.add_paragraph('• Férias proporcionais + 1/3: R$ [VALOR]')
        doc_recibo.add_paragraph('• 13º salário proporcional: R$ [VALOR]')
        doc_recibo.add_paragraph('• Aviso prévio: R$ [VALOR]')
        doc_recibo.add_paragraph('• Outras verbas: R$ [VALOR]')
        
        doc_recibo.add_paragraph('\\nTotal líquido recebido: R$ [TOTAL]')
        
        doc_recibo.add_paragraph(f'\\n{dados_usuario.get("cidade", "[CIDADE]")}, {datetime.now().strftime("%d de %B de %Y")}')
        doc_recibo.add_paragraph('\\n\\n_________________________________')
        doc_recibo.add_paragraph(f'{dados_usuario.get("nome", "[NOME]")}')
        
        # Salva em buffer
        buffer_recibo = BytesIO()
        doc_recibo.save(buffer_recibo)
        buffer_recibo.seek(0)
        documentos['recibo_quitacao.docx'] = buffer_recibo
        
        # 2. Requerimento FGTS
        doc_fgts = Document()
        doc_fgts.add_heading('REQUERIMENTO DE SAQUE DO FGTS', 0)
        
        doc_fgts.add_paragraph('À Caixa Econômica Federal')
        doc_fgts.add_paragraph('\\nEu, ' + dados_usuario.get('nome', '[NOME]') + 
                              ', portador do CPF nº [CPF], venho requerer o saque do FGTS ' +
                              'referente ao término do contrato de trabalho.')
        
        doc_fgts.add_paragraph('\\nDados do contrato:')
        doc_fgts.add_paragraph(f'• Empresa: {dados_usuario.get("empresa", "[EMPRESA]")}')
        doc_fgts.add_paragraph(f'• Período: {dados_usuario.get("data_admissao", "[DATA_ADMISSAO]")} a {dados_usuario.get("data_demissao", "[DATA_DEMISSAO]")}')
        doc_fgts.add_paragraph('• Motivo: Rescisão sem justa causa')
        
        doc_fgts.add_paragraph('\\nSolicito a liberação dos valores depositados na conta vinculada.')
        
        doc_fgts.add_paragraph(f'\\n{dados_usuario.get("cidade", "[CIDADE]")}, {datetime.now().strftime("%d de %B de %Y")}')
        doc_fgts.add_paragraph('\\n\\n_________________________________')
        doc_fgts.add_paragraph(f'{dados_usuario.get("nome", "[NOME]")}')
        
        buffer_fgts = BytesIO()
        doc_fgts.save(buffer_fgts)
        buffer_fgts.seek(0)
        documentos['requerimento_fgts.docx'] = buffer_fgts
        
        # 3. Carta de Negociação
        doc_negociacao = Document()
        doc_negociacao.add_heading('PROPOSTA DE ACORDO TRABALHISTA', 0)
        
        doc_negociacao.add_paragraph(f'À {dados_usuario.get("empresa", "[EMPRESA]")}')
        doc_negociacao.add_paragraph('A/C Departamento de Recursos Humanos')
        
        doc_negociacao.add_paragraph('\\nVenho por meio desta apresentar proposta de acordo para rescisão do contrato de trabalho.')
        
        doc_negociacao.add_paragraph('\\nProposta:')
        doc_negociacao.add_paragraph('• Rescisão por acordo (Art. 484-A da CLT)')
        doc_negociacao.add_paragraph('• Pagamento de 50% do aviso prévio')
        doc_negociacao.add_paragraph('• Pagamento de 20% da multa do FGTS')
        doc_negociacao.add_paragraph('• Demais verbas integrais')
        
        doc_negociacao.add_paragraph('\\nAguardo retorno para agendamento da rescisão.')
        
        doc_negociacao.add_paragraph(f'\\n{dados_usuario.get("cidade", "[CIDADE]")}, {datetime.now().strftime("%d de %B de %Y")}')
        doc_negociacao.add_paragraph('\\n\\nAtenciosamente,')
        doc_negociacao.add_paragraph('\\n_________________________________')
        doc_negociacao.add_paragraph(f'{dados_usuario.get("nome", "[NOME]")}')
        
        buffer_negociacao = BytesIO()
        doc_negociacao.save(buffer_negociacao)
        buffer_negociacao.seek(0)
        documentos['carta_negociacao.docx'] = buffer_negociacao
        
        return documentos
    
    def gerar_planilha_xlsx(self, dados_calculo, dados_usuario):
        """Gera planilha financeira com fórmulas"""
        wb = Workbook()
        
        # Aba 1: Controle de Verbas
        ws1 = wb.active
        ws1.title = "Controle de Verbas"
        
        # Cabeçalho
        ws1['A1'] = 'CONTROLE DE VERBAS RESCISÓRIAS'
        ws1['A1'].font = Font(size=16, bold=True)
        ws1.merge_cells('A1:D1')
        
        # Dados
        row = 3
        ws1[f'A{row}'] = 'Verba'
        ws1[f'B{row}'] = 'Valor Calculado'
        ws1[f'C{row}'] = 'Valor Recebido'
        ws1[f'D{row}'] = 'Diferença'
        
        # Estilo do cabeçalho
        for col in ['A', 'B', 'C', 'D']:
            cell = ws1[f'{col}{row}']
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color='3498DB', end_color='3498DB', fill_type='solid')
        
        # Dados das verbas
        verbas_data = [
            ('Saldo de Salário', dados_calculo['verbas']['saldo_salario']['valor']),
            ('Férias Proporcionais', dados_calculo['verbas']['ferias_proporcionais']['total']),
            ('13º Proporcional', dados_calculo['verbas']['decimo_terceiro']['valor']),
            ('Aviso Prévio', dados_calculo['verbas']['aviso_previo']['valor']),
            ('Multa FGTS 40%', dados_calculo['verbas']['multa_fgts']['valor']),
        ]
        
        for i, (nome, valor) in enumerate(verbas_data, start=row+1):
            ws1[f'A{i}'] = nome
            ws1[f'B{i}'] = float(valor)
            ws1[f'C{i}'] = 0  # Para o usuário preencher
            ws1[f'D{i}'] = f'=C{i}-B{i}'  # Fórmula da diferença
        
        # Aba 2: Planejamento Financeiro
        ws2 = wb.create_sheet("Planejamento")
        
        ws2['A1'] = 'PLANEJAMENTO FINANCEIRO'
        ws2['A1'].font = Font(size=16, bold=True)
        ws2.merge_cells('A1:C1')
        
        # Categorias de gastos
        categorias = [
            'Reserva de Emergência (6 meses)',
            'Quitação de Dívidas',
            'Investimentos',
            'Gastos Pessoais',
            'Capacitação/Cursos'
        ]
        
        row = 3
        ws2[f'A{row}'] = 'Categoria'
        ws2[f'B{row}'] = 'Valor Planejado'
        ws2[f'C{row}'] = 'Percentual'
        
        total_liquido = float(dados_calculo['resumo']['total_liquido'])
        
        for i, categoria in enumerate(categorias, start=row+1):
            ws2[f'A{i}'] = categoria
            ws2[f'B{i}'] = 0  # Para o usuário preencher
            ws2[f'C{i}'] = f'=B{i}/{total_liquido}'  # Percentual
        
        # Aba 3: Simulador de Investimentos
        ws3 = wb.create_sheet("Investimentos")
        
        ws3['A1'] = 'SIMULADOR DE INVESTIMENTOS'
        ws3['A1'].font = Font(size=16, bold=True)
        ws3.merge_cells('A1:E1')
        
        # Simulação
        ws3['A3'] = 'Valor Inicial:'
        ws3['B3'] = total_liquido
        ws3['A4'] = 'Taxa Mensal (%):'
        ws3['B4'] = 1.0  # 1% ao mês
        ws3['A5'] = 'Período (meses):'
        ws3['B5'] = 12
        
        ws3['A7'] = 'Valor Final:'
        ws3['B7'] = '=B3*(1+B4/100)^B5'
        
        ws3['A8'] = 'Rendimento:'
        ws3['B8'] = '=B7-B3'
        
        # Formatação
        for ws in [ws1, ws2, ws3]:
            for column in ws.columns:
                max_length = 0
                column_letter = get_column_letter(column[0].column)
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = min(max_length + 2, 50)
                ws.column_dimensions[column_letter].width = adjusted_width
        
        # Salva em buffer
        buffer = BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        
        return buffer
    
    def gerar_guia_pdf(self):
        """Gera guia estratégico em PDF"""
        
        html_content = """
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <style>
                body { font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; }
                .header { text-align: center; margin-bottom: 30px; }
                .logo { font-size: 24px; font-weight: bold; color: #2c3e50; }
                .section { margin: 20px 0; }
                .section h2 { color: #2c3e50; border-bottom: 2px solid #3498db; padding-bottom: 5px; }
                .checklist { background-color: #f8f9fa; padding: 15px; border-left: 4px solid #28a745; }
                .tip { background-color: #e8f5e8; padding: 10px; border-radius: 5px; margin: 10px 0; }
            </style>
        </head>
        <body>
            <div class="header">
                <div class="logo">GUIA ESTRATÉGICO DE RESCISÃO</div>
                <div>Como maximizar seus direitos trabalhistas</div>
            </div>
            
            <div class="section">
                <h2>🎯 Como Negociar sua Demissão</h2>
                <p><strong>1. Prepare-se antes da conversa:</strong></p>
                <ul>
                    <li>Calcule suas verbas rescisórias</li>
                    <li>Conheça seus direitos</li>
                    <li>Tenha documentos organizados</li>
                </ul>
                
                <p><strong>2. Estratégias de negociação:</strong></p>
                <ul>
                    <li>Proponha rescisão por acordo (Art. 484-A)</li>
                    <li>Negocie data de saída favorável</li>
                    <li>Solicite carta de recomendação</li>
                </ul>
            </div>
            
            <div class="section">
                <h2>📋 Checklist Pós-Demissão</h2>
                <div class="checklist">
                    <p><strong>Primeiros 7 dias:</strong></p>
                    <ul>
                        <li>☐ Conferir cálculo das verbas</li>
                        <li>☐ Solicitar documentos (TRCT, CD/SD)</li>
                        <li>☐ Dar entrada no seguro-desemprego</li>
                        <li>☐ Sacar FGTS</li>
                    </ul>
                    
                    <p><strong>Primeiros 30 dias:</strong></p>
                    <ul>
                        <li>☐ Atualizar currículo</li>
                        <li>☐ Planejar uso da rescisão</li>
                        <li>☐ Buscar recolocação</li>
                        <li>☐ Considerar capacitação</li>
                    </ul>
                </div>
            </div>
            
            <div class="section">
                <h2>💡 Dicas Importantes</h2>
                <div class="tip">
                    <strong>Dica 1:</strong> Sempre confira os cálculos antes de assinar a rescisão.
                </div>
                <div class="tip">
                    <strong>Dica 2:</strong> Guarde todos os documentos por pelo menos 5 anos.
                </div>
                <div class="tip">
                    <strong>Dica 3:</strong> Em caso de dúvidas, procure o sindicato da categoria.
                </div>
            </div>
        </body>
        </html>
        """
        
        pdf_buffer = BytesIO()
        HTML(string=html_content).write_pdf(pdf_buffer)
        pdf_buffer.seek(0)
        
        return pdf_buffer
    
    def gerar_kit_completo(self, dados_calculo, dados_usuario):
        """Gera o kit completo e retorna arquivo ZIP"""
        
        # Cria ZIP em memória
        zip_buffer = BytesIO()
        
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            
            # 1. Relatório PDF
            relatorio_pdf = self.gerar_relatorio_pdf(dados_calculo, dados_usuario)
            zip_file.writestr('01_Relatorio_Personalizado.pdf', relatorio_pdf.read())
            
            # 2. Modelos DOCX
            modelos_docx = self.gerar_modelos_docx(dados_usuario)
            for nome, buffer in modelos_docx.items():
                zip_file.writestr(f'02_Modelos/{nome}', buffer.read())
            
            # 3. Planilha XLSX
            planilha_xlsx = self.gerar_planilha_xlsx(dados_calculo, dados_usuario)
            zip_file.writestr('03_Planilha_Financeira.xlsx', planilha_xlsx.read())
            
            # 4. Guia PDF
            guia_pdf = self.gerar_guia_pdf()
            zip_file.writestr('04_Guia_Estrategico.pdf', guia_pdf.read())
        
        zip_buffer.seek(0)
        return zip_buffer
    
    def _formatar_tipo_demissao(self, tipo):
        """Formata o tipo de demissão para exibição"""
        tipos = {
            'sem_justa_causa': 'Demissão sem justa causa',
            'justa_causa': 'Demissão por justa causa',
            'pedido_demissao': 'Pedido de demissão'
        }
        return tipos.get(tipo, tipo)
    
    def _gerar_linhas_verbas(self, dados_calculo):
        """Gera as linhas da tabela de verbas"""
        verbas = dados_calculo['verbas']
        linhas = []
        
        # Saldo de salário
        saldo = verbas['saldo_salario']
        linhas.append(f'<tr><td>Saldo de Salário</td><td>{saldo["dias_trabalhados"]} dias trabalhados</td><td>R$ {saldo["valor"]:,.2f}</td></tr>')
        
        # Férias
        ferias = verbas['ferias_proporcionais']
        linhas.append(f'<tr><td>Férias Proporcionais</td><td>{ferias["meses_trabalhados"]} meses</td><td>R$ {ferias["valor_ferias"]:,.2f}</td></tr>')
        linhas.append(f'<tr><td>1/3 Constitucional</td><td>1/3 das férias</td><td>R$ {ferias["um_terco"]:,.2f}</td></tr>')
        
        # 13º
        decimo = verbas['decimo_terceiro']
        linhas.append(f'<tr><td>13º Proporcional</td><td>{decimo["meses_trabalhados"]} meses</td><td>R$ {decimo["valor"]:,.2f}</td></tr>')
        
        # Aviso prévio
        aviso = verbas['aviso_previo']
        if aviso['valor'] > 0:
            linhas.append(f'<tr><td>Aviso Prévio</td><td>{aviso["dias"]} dias ({aviso["tipo"]})</td><td>R$ {aviso["valor"]:,.2f}</td></tr>')
        
        # Multa FGTS
        multa = verbas['multa_fgts']
        if multa['aplicavel']:
            linhas.append(f'<tr><td>Multa FGTS</td><td>{multa["percentual"]}% do saldo</td><td>R$ {multa["valor"]:,.2f}</td></tr>')
        
        return '\\n'.join(linhas)
    
    def _gerar_recomendacoes(self, dados_calculo, dados_usuario):
        """Gera recomendações personalizadas"""
        total_liquido = dados_calculo['resumo']['total_liquido']
        
        recomendacoes = []
        
        if total_liquido > 10000:
            recomendacoes.append('<p>💰 <strong>Alto valor de rescisão:</strong> Considere investir parte em aplicações de longo prazo.</p>')
        
        if dados_usuario.get('tipo_demissao') == 'sem_justa_causa':
            recomendacoes.append('<p>🎯 <strong>Seguro-desemprego:</strong> Você tem direito ao benefício. Solicite nos primeiros 120 dias.</p>')
        
        recomendacoes.append('<p>📚 <strong>Capacitação:</strong> Use parte dos recursos para cursos que aumentem sua empregabilidade.</p>')
        recomendacoes.append('<p>🏦 <strong>Reserva de emergência:</strong> Mantenha pelo menos 6 meses de gastos em aplicação líquida.</p>')
        
        return '\\n'.join(recomendacoes)

